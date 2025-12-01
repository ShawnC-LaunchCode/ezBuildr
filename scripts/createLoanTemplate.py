#!/usr/bin/env python3
"""
Create a professional loan application summary template
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ==========================
# HEADER / LETTERHEAD
# ==========================
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = header.add_run('FINANCIAL SERVICES COMPANY')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

header2 = doc.add_paragraph()
header2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run2 = header2.add_run('123 Main Street • Suite 500 • Anytown, ST 12345\n')
run2.font.size = Pt(9)
run3 = header2.add_run('Phone: (555) 123-4567 • Email: loans@financialservices.com')
run3.font.size = Pt(9)

doc.add_paragraph('_' * 80)  # Divider line

# ==========================
# DATE AND REFERENCE
# ==========================
doc.add_paragraph()
date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_paragraph()

# ==========================
# APPLICANT INFORMATION
# ==========================
title = doc.add_heading('LOAN APPLICATION SUMMARY', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

# Personal Information Section
doc.add_heading('Applicant Information', level=2)

applicant_table = doc.add_table(rows=5, cols=2)
applicant_table.style = 'Light Grid Accent 1'

applicant_table.rows[0].cells[0].text = 'Full Name:'
applicant_table.rows[0].cells[1].text = '{{firstName}} {{lastName}}'

applicant_table.rows[1].cells[0].text = 'Email Address:'
applicant_table.rows[1].cells[1].text = '{{email}}'

applicant_table.rows[2].cells[0].text = 'Phone Number:'
applicant_table.rows[2].cells[1].text = '{{phone}}'

applicant_table.rows[3].cells[0].text = 'Date of Birth:'
applicant_table.rows[3].cells[1].text = '{{dateOfBirth}}'

applicant_table.rows[4].cells[0].text = 'SSN (last 4 digits):'
applicant_table.rows[4].cells[1].text = 'XXX-XX-{{ssn}}'

doc.add_paragraph()

# ==========================
# EMPLOYMENT INFORMATION
# ==========================
doc.add_heading('Employment Information', level=2)

employment_table = doc.add_table(rows=4, cols=2)
employment_table.style = 'Light Grid Accent 1'

employment_table.rows[0].cells[0].text = 'Employment Status:'
employment_table.rows[0].cells[1].text = '{{employmentStatus}}'

employment_table.rows[1].cells[0].text = 'Employer Name:'
employment_table.rows[1].cells[1].text = '{{employerName}}'

employment_table.rows[2].cells[0].text = 'Job Title:'
employment_table.rows[2].cells[1].text = '{{jobTitle}}'

employment_table.rows[3].cells[0].text = 'Annual Income:'
employment_table.rows[3].cells[1].text = '${{annualIncome}}'

doc.add_paragraph()

# ==========================
# LOAN REQUEST DETAILS
# ==========================
doc.add_heading('Loan Request Details', level=2)

loan_table = doc.add_table(rows=3, cols=2)
loan_table.style = 'Light Grid Accent 1'

loan_table.rows[0].cells[0].text = 'Requested Amount:'
run = loan_table.rows[0].cells[1].paragraphs[0].add_run('${{loanAmount}}')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 102, 0)

loan_table.rows[1].cells[0].text = 'Loan Purpose:'
loan_table.rows[1].cells[1].text = '{{loanPurpose}}'

loan_table.rows[2].cells[0].text = 'Preferred Term:'
loan_table.rows[2].cells[1].text = '{{loanTerm}}'

doc.add_paragraph()

# ==========================
# FINANCIAL ANALYSIS
# ==========================
doc.add_heading('Financial Analysis', level=2)

analysis_table = doc.add_table(rows=3, cols=2)
analysis_table.style = 'Light Grid Accent 1'

analysis_table.rows[0].cells[0].text = 'Monthly Income:'
analysis_table.rows[0].cells[1].text = '${{debtToIncomeRatio.monthlyIncome}}'

analysis_table.rows[1].cells[0].text = 'Monthly Debt Payments:'
analysis_table.rows[1].cells[1].text = '${{monthlyDebt}}'

analysis_table.rows[2].cells[0].text = 'Debt-to-Income Ratio:'
analysis_table.rows[2].cells[1].text = '{{debtToIncomeRatio.ratio}}% ({{debtToIncomeRatio.status}})'

doc.add_paragraph()

# ==========================
# FOOTER / DISCLAIMER
# ==========================
doc.add_paragraph('_' * 80)  # Divider line

doc.add_paragraph()

disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = disclaimer.add_run('APPLICATION STATUS: PENDING REVIEW')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

next_steps = doc.add_paragraph()
next_steps.add_run('Next Steps:\n').bold = True
next_steps.add_run(
    '1. A loan officer will review your application within 2-3 business days\n'
    '2. You may be contacted for additional documentation\n'
    '3. Upon approval, you will receive loan terms and documentation to sign\n'
)

doc.add_paragraph()

footer_para = doc.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_run = footer_para.add_run(
    'This document is confidential and intended solely for the named applicant.\n'
    'Please contact us at (555) 123-4567 with any questions.'
)
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# ==========================
# SAVE DOCUMENT
# ==========================
output_dir = os.path.join(os.path.dirname(__file__), '..', 'server', 'files', 'templates')
os.makedirs(output_dir, exist_ok=True)

output_path = os.path.join(output_dir, 'loan-application-summary.docx')
doc.save(output_path)

print(f'[OK] Template created: {output_path}')
print('\nTemplate includes the following placeholders:')
print('  - {{firstName}}, {{lastName}}')
print('  - {{email}}, {{phone}}, {{dateOfBirth}}, {{ssn}}')
print('  - {{employmentStatus}}, {{employerName}}, {{jobTitle}}')
print('  - {{annualIncome}}, {{monthlyDebt}}')
print('  - {{loanAmount}}, {{loanPurpose}}, {{loanTerm}}')
print('  - {{debtToIncomeRatio.ratio}}, {{debtToIncomeRatio.status}}, {{debtToIncomeRatio.monthlyIncome}}')
print('\nTemplate ready for use!')
