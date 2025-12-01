#!/usr/bin/env python3
"""
Create a simple DOCX template for testing document generation
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('Welcome Letter', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add some paragraphs with placeholders
doc.add_paragraph()
doc.add_paragraph('Dear {{firstName}} {{lastName}},')
doc.add_paragraph()

body = doc.add_paragraph(
    'Thank you for completing our workflow! We have successfully received your information.'
)

doc.add_paragraph()
doc.add_paragraph('Your Details:')

# Add a table with user details
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
table.rows[0].cells[0].text = 'First Name:'
table.rows[0].cells[1].text = '{{firstName}}'

table.rows[1].cells[0].text = 'Last Name:'
table.rows[1].cells[1].text = '{{lastName}}'

table.rows[2].cells[0].text = 'Email:'
table.rows[2].cells[1].text = '{{email}}'

doc.add_paragraph()
doc.add_paragraph('We look forward to working with you!')

doc.add_paragraph()
closing = doc.add_paragraph('Best regards,')
doc.add_paragraph('The VaultLogic Team')

# Save the document
output_path = os.path.join(os.path.dirname(__file__), '..', 'uploads', 'templates', 'welcome-letter.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)

print(f'Template created: {output_path}')
print('Template includes placeholders: {{firstName}}, {{lastName}}, {{email}}')
